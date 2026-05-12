import csv
import os
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CSV_FILE = 'FICHA DE INSCRIÇÃO - V EJC NAZARÉ (respostas) - Página3.csv'
OUTPUT_DIR = 'Fichas de Inscrição'

# ── Helpers de formatação ──────────────────────────────────────────────────────

def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '000000')
        tcBorders.append(border)
    tcPr.append(tcBorders)

def set_cell_bg(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def set_row_height(row, height_cm):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_cm * 567)))
    trHeight.set(qn('w:hRule'), 'exact')
    trPr.append(trHeight)

def cell_text(cell, text, bold=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text) if text else '')
    run.bold = bold
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def label(cell, txt, bg='D9D9D9'):
    cell_text(cell, txt, bold=True, size=9)
    set_cell_bg(cell, bg)

def value(cell, txt, size=11):
    cell_text(cell, txt or '', bold=False, size=size)

def checkboxes(cell, options, selected, size=10):
    """Renderiza opções com [X] para selecionada(s) e [ ] para as demais."""
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    selected_lower = [s.strip().lower() for s in (selected or '').split(';')]

    for i, opt in enumerate(options):
        marked = any(opt.lower() in s or s in opt.lower() for s in selected_lower if s)
        prefix = '[X] ' if marked else '[  ] '
        run = p.add_run(prefix + opt)
        run.font.size = Pt(size)
        run.bold = marked
        if i < len(options) - 1:
            p.add_run('   ').font.size = Pt(size)

# ── Geração da ficha ──────────────────────────────────────────────────────────

def generate_ficha(data: dict, output_path: str):
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    # ── Cabeçalho ─────────────────────────────────────────────
    ht = doc.add_table(rows=1, cols=2)
    ht.alignment = WD_TABLE_ALIGNMENT.CENTER
    hr = ht.rows[0]
    set_row_height(hr, 1.8)

    lc = hr.cells[0]
    lc.width = Cm(14)
    p = lc.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Encontro de Jovens com Cristo\nParóquia Nossa Senhora de Nazaré')
    r.bold = True
    r.font.size = Pt(16)

    rc = hr.cells[1]
    rc.width = Cm(4)
    p2 = rc.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run('FOTO\n3x4')
    r2.font.size = Pt(12)
    set_cell_border(rc)

    title = doc.add_paragraph('Ficha de Inscrição')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(14)
    title.paragraph_format.space_before = Pt(4)
    title.paragraph_format.space_after = Pt(2)

    # ── Tabela principal ──────────────────────────────────────
    t = doc.add_table(rows=0, cols=6)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    def row(h=0.55):
        r = t.add_row()
        set_row_height(r, h)
        for c in r.cells:
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        return r.cells

    # Linha 1 — Nº Ficha | Data Nasc | RG | CPF
    c = row()
    c[0].merge(c[1]); label(c[0], 'Nº FICHA:')
    c[2].merge(c[3]); label(c[2], 'DATA DE NASCIMENTO:')
    label(c[4], 'RG:')
    label(c[5], 'CPF:')

    c = row(0.6)
    c[0].merge(c[1])
    c[2].merge(c[3]); value(c[2], data['nasc'])
    value(c[4], data['rg'])
    value(c[5], data['cpf'])

    # Linha 2 — Nome
    c = row()
    c[0].merge(c[5]); label(c[0], 'NOME COMPLETO:')
    c = row(0.6)
    c[0].merge(c[5]); value(c[0], data['nome'])

    # Linha 3 — Endereço
    c = row()
    c[0].merge(c[5]); label(c[0], 'ENDEREÇO COMPLETO:')
    c = row(0.6)
    c[0].merge(c[5]); value(c[0], data['endereco'])

    # Linha 4 — Bairro | Telefone | Ponto de Referência
    c = row()
    c[0].merge(c[1]); label(c[0], 'BAIRRO:')
    c[2].merge(c[3]); label(c[2], 'TELEFONE (DDD + nº):')
    c[4].merge(c[5]); label(c[4], 'PONTO DE REFERÊNCIA:')
    c = row(0.6)
    c[0].merge(c[1]); value(c[0], data['bairro'])
    c[2].merge(c[3]); value(c[2], data['tel'])
    c[4].merge(c[5]); value(c[4], data['ref'])

    # Linha 5 — E-mail | Estado Civil
    c = row()
    c[0].merge(c[2]); label(c[0], 'E-MAIL:')
    c[3].merge(c[5]); label(c[3], 'ESTADO CIVIL:')
    c = row(0.65)
    c[0].merge(c[2]); value(c[0], data['email'])
    c[3].merge(c[5])
    checkboxes(c[3], ['Solteiro', 'Casado', 'Separado', 'Mora junto', 'Outro'], data['estado_civil'])

    # Linha 6 — Como quer ser chamado | Religião
    c = row()
    c[0].merge(c[2]); label(c[0], 'COMO QUER SER CHAMADO NO ENCONTRO:')
    c[3].merge(c[5]); label(c[3], 'RELIGIÃO:')
    c = row(0.6)
    c[0].merge(c[2]); value(c[0], data['apelido'])
    c[3].merge(c[5]); value(c[3], data['religiao'])

    # Linha 7 — Necessidade especial
    c = row()
    c[0].merge(c[1]); label(c[0], 'PORTADOR DE NECESSIDADE ESPECIAL?')
    c[2].merge(c[3])
    checkboxes(c[2], ['SIM', 'NÃO'], data['nec_especial'])
    c[4].merge(c[5]); label(c[4], 'SE SIM, QUAL:')
    c = row(0.6)
    c[4].merge(c[5]); value(c[4], data['nec_especial_qual'])

    # Linha 8 — Mora com os pais | Tem filhos
    c = row()
    c[0].merge(c[1]); label(c[0], 'MORA COM OS PAIS?')
    c[2].merge(c[3])
    checkboxes(c[2], ['SIM', 'NÃO'], data['mora_pais'])
    c[4].merge(c[5]); label(c[4], 'TEM FILHO(S)?')
    c = row(0.55)
    c[0].merge(c[3])
    c[4].merge(c[5])
    checkboxes(c[4], ['SIM', 'NÃO'], data['filhos'])

    # Linha 9 — Mãe
    c = row()
    c[0].merge(c[3]); label(c[0], 'NOME DA MÃE:')
    c[4].merge(c[5]); label(c[4], 'TELEFONE DA MÃE:')
    c = row(0.6)
    c[0].merge(c[3]); value(c[0], data['mae'])
    c[4].merge(c[5]); value(c[4], data['tel_mae'])

    # Linha 10 — Pai
    c = row()
    c[0].merge(c[3]); label(c[0], 'NOME DO PAI:')
    c[4].merge(c[5]); label(c[4], 'TELEFONE DO PAI:')
    c = row(0.6)
    c[0].merge(c[3]); value(c[0], data['pai'])
    c[4].merge(c[5]); value(c[4], data['tel_pai'])

    # Linha 11 — Irmão
    c = row()
    c[0].merge(c[3]); label(c[0], 'NOME DE UM IRMÃO (caso tenha):')
    c[4].merge(c[5]); label(c[4], 'TELEFONE DO IRMÃO:')
    c = row(0.6)
    c[0].merge(c[3]); value(c[0], data['irmao'])
    c[4].merge(c[5]); value(c[4], data['tel_irmao'])

    # Linha 12 — Amigo próximo
    c = row()
    c[0].merge(c[3]); label(c[0], 'NOME DE UM AMIGO PRÓXIMO:')
    c[4].merge(c[5]); label(c[4], 'TELEFONE DO AMIGO:')
    c = row(0.6)
    c[0].merge(c[3]); value(c[0], data['amigo'])
    c[4].merge(c[5]); value(c[4], data['tel_amigo'])

    # Linha 13 — Pais na igreja
    c = row()
    c[0].merge(c[5]); label(c[0], 'SEUS PAIS FAZEM PARTE DE ALGUMA IGREJA? SE SIM, QUAL?')
    c = row(0.6)
    c[0].merge(c[5]); value(c[0], data['pais_igreja'])

    # Linha 14 — Sacramentos | Igreja que frequenta
    c = row()
    c[0].merge(c[2]); label(c[0], 'SACRAMENTOS:')
    c[3].merge(c[5]); label(c[3], 'QUAL IGREJA FREQUENTA?')
    c = row(0.65)
    c[0].merge(c[2])
    checkboxes(c[0], ['Batismo', '1ª Eucaristia', 'Crisma'], data['sacramentos'])
    c[3].merge(c[5]); value(c[3], data['igreja'])

    # Linha 15 — Movimento/Pastoral
    c = row()
    c[0].merge(c[1]); label(c[0], 'FAZ PARTE DE ALGUM MOVIMENTO OU PASTORAL?')
    c[2].merge(c[3])
    checkboxes(c[2], ['SIM', 'NÃO'], data['pastoral'])
    c[4].merge(c[5]); label(c[4], 'SE SIM, QUAL?')
    c = row(0.6)
    c[4].merge(c[5]); value(c[4], data['pastoral_qual'])

    # Linha 16 — Alergia
    c = row()
    c[0].merge(c[1]); label(c[0], 'ALERGIA A REMÉDIO OU COMIDA?')
    c[2].merge(c[3])
    checkboxes(c[2], ['SIM', 'NÃO'], data['alergia'])
    c[4].merge(c[5]); label(c[4], 'SE SIM, QUAL(AIS)?')
    c = row(0.6)
    c[4].merge(c[5]); value(c[4], data['alergia_qual'])

    # Linha 17 — Remédio diário
    c = row()
    c[0].merge(c[1]); label(c[0], 'TOMA ALGUM REMÉDIO DIARIAMENTE?')
    c[2].merge(c[3])
    checkboxes(c[2], ['SIM', 'NÃO'], data['remedio'])
    c[4].merge(c[5]); label(c[4], 'SE SIM, QUAL(AIS) E HORÁRIOS?')
    c = row(0.7)
    c[4].merge(c[5]); value(c[4], f"{data['remedio_qual']}  {data['remedio_horario']}".strip())

    # Linha 18 — Relacionamentos
    rels = [
        ('RELACIONAMENTO COM OS PAIS:', data['rel_pais']),
        ('RELACIONAMENTO COM OS IRMÃOS:', data['rel_irmaos']),
        ('RELACIONAMENTO COM OS AMIGOS:', data['rel_amigos']),
    ]
    opts_rel = ['Aberto/Franco', 'Moderado', 'Tímido', 'Não há diálogo']
    for lbl_txt, val_txt in rels:
        c = row(0.7)
        c[0].merge(c[2]); label(c[0], lbl_txt)
        c[3].merge(c[5])
        checkboxes(c[3], opts_rel, val_txt)

    # Linha 19 — Quem convidou
    c = row()
    c[0].merge(c[5]); label(c[0], 'ALGUÉM TE CONVIDOU PARA PARTICIPAR? SE SIM, QUEM?')
    c = row(0.6)
    c[0].merge(c[5]); value(c[0], data['convidou'])

    # Linha 20 — Como soube
    c = row()
    c[0].merge(c[5]); label(c[0], 'COMO FICOU SABENDO DAS INSCRIÇÕES?')
    c = row(0.65)
    c[0].merge(c[5])
    checkboxes(c[0], ['Rede Social', 'Avisos na Missa', 'Amigos/Familiares que já participaram', 'Outros'], data['como_soube'])

    # Linha 21 — Conhece alguém que vai
    c = row()
    c[0].merge(c[5]); label(c[0], 'CONHECE ALGUÉM QUE TAMBÉM VAI PARTICIPAR? SE SIM, QUEM?')
    c = row(0.6)
    c[0].merge(c[5]); value(c[0], data['conhece_alguem'])

    # Linha 22 — Por que deseja
    c = row()
    c[0].merge(c[5]); label(c[0], 'POR QUE DESEJA FAZER ESTE ENCONTRO E O QUE ESPERA?')
    c = row(1.4)
    c[0].merge(c[5]); value(c[0], data['por_que'], size=10)

    # Linha 23 — Tamanho camisa | Instagram | Observações
    c = row()
    c[0].merge(c[1]); label(c[0], 'TAMANHO DA CAMISA:')
    c[2].merge(c[3]); label(c[2], 'INSTAGRAM:')
    c[4].merge(c[5]); label(c[4], 'OBSERVAÇÕES:')
    c = row(0.65)
    c[0].merge(c[1])
    checkboxes(c[0], ['PP', 'P', 'M', 'G', 'GG'], data['camisa'], size=10)
    c[2].merge(c[3]); value(c[2], data['instagram'])
    c[4].merge(c[5]); value(c[4], data['obs'])

    # Linha 24 — OBS Grupo Dirigente
    c = row()
    c[0].merge(c[5]); label(c[0], 'OBS DO GRUPO DIRIGENTE:   ( ) TAXA   ( ) FOTO')
    c = row(0.8)
    c[0].merge(c[5])

    # Assinaturas
    c = row(1.6)
    c[0].merge(c[2])
    cell_text(c[0], '\n___________________________________\nASSINATURA GRUPO DIRIGENTE', size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    c[3].merge(c[5])
    cell_text(c[3], '\n___________________________________\nASSINATURA DO ENCONTRISTA', size=11, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.save(output_path)

# ── Leitura do CSV ────────────────────────────────────────────────────────────

def parse_row(row: list) -> dict:
    def g(i):
        return row[i].strip() if i < len(row) else ''

    # Para campos SIM/NÃO, o Forms retorna 'Sim' ou 'Não'
    def sim_nao(i):
        v = g(i).lower()
        if 'sim' in v: return 'SIM'
        if 'não' in v or 'nao' in v: return 'NÃO'
        return g(i)

    # Estado civil: pega só a primeira palavra capitalizada
    ec = g(9).strip()

    return {
        'nome':             g(1),
        'nasc':             g(2),
        'rg':               g(3),
        'cpf':              g(4),
        'endereco':         g(5),
        'bairro':           g(6),
        'ref':              g(7),
        'tel':              g(8),
        'estado_civil':     ec,
        'email':            g(10),
        'nec_especial':     sim_nao(11),
        'nec_especial_qual':g(12),
        'apelido':          g(13),
        'religiao':         g(14),
        'sacramentos':      g(15),   # pode ser "Batismo; Crisma" etc.
        'igreja':           g(16),
        'pastoral':         g(17),   # pode ser "Não" ou nome do movimento
        'pastoral_qual':    g(17) if g(17).lower() not in ('sim','não','nao','') else '',
        'mora_pais':        sim_nao(18),
        'filhos':           sim_nao(19),
        'mae':              g(20),
        'tel_mae':          g(21),
        'pai':              g(22),
        'tel_pai':          g(23),
        'irmao':            g(24),
        'tel_irmao':        g(25),
        'amigo':            g(26),
        'tel_amigo':        g(27),
        'pais_igreja':      g(28),
        'alergia':          sim_nao(29),
        'alergia_qual':     g(29) if g(29).lower() not in ('sim','não','nao','') else '',
        'remedio':          sim_nao(30),
        'remedio_qual':     g(30) if g(30).lower() not in ('sim','não','nao','') else '',
        'remedio_horario':  g(31),
        'rel_pais':         g(32),
        'rel_irmaos':       g(33),
        'rel_amigos':       g(34),
        'convidou':         g(35),
        'como_soube':       g(36),
        'por_que':          g(37),
        'conhece_alguem':   g(38),
        'camisa':           g(39),
        'obs':              g(40),
        'instagram':        g(43) if len(row) > 43 else '',
    }

def safe_filename(name: str) -> str:
    return re.sub(r'[<>:"/\\|?*]', '', name).strip()

# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    encodings = ['utf-8-sig', 'latin-1', 'cp1252']
    rows = []
    used_enc = None
    for enc in encodings:
        try:
            with open(CSV_FILE, encoding=enc, newline='') as f:
                reader = csv.reader(f)
                rows = list(reader)
            used_enc = enc
            break
        except Exception:
            continue

    if not rows:
        print('Erro: não foi possível ler o CSV.')
        return

    print(f'Lendo com encoding {used_enc} — {len(rows)} fichas encontradas.\n')

    for i, row_data in enumerate(rows, 1):
        data = parse_row(row_data)
        nome = data['nome'] or f'Sem_Nome_{i}'
        filename = f'FICHA DE INSCRIÇÃO - {safe_filename(nome)}.docx'
        output_path = os.path.join(OUTPUT_DIR, filename)
        generate_ficha(data, output_path)
        print(f'[{i:02d}] {filename}')

    print(f'\nOK: {len(rows)} fichas geradas em "{OUTPUT_DIR}/"')

if __name__ == '__main__':
    main()